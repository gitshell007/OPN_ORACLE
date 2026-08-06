"""Editable DOCX export for durable opportunity offer drafts (SV2-G09-B).

Uses ``python-docx`` (pinned in pyproject/uv.lock) so Word/LibreOffice open the
file without repair. Content always comes from the persisted
``OpportunityOfferDraft.content`` snapshot (never regenerated from artifacts).
"""

from __future__ import annotations

import io
import re
import uuid
from collections.abc import Mapping
from dataclasses import dataclass
from datetime import UTC, datetime
from typing import Any
from urllib.parse import quote
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from opn_oracle.ai.draft_offer import _BANNER, _HUMAN_GATE, _RESPONSE_TAG
from opn_oracle.ai.offer_draft import MAX_CONTENT_BYTES, OfferDraftError

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_DOCX_BYTES = 2_000_000
MAX_DOCX_PARAGRAPHS = 2_000
MAX_TRACEABILITY_ROWS = 60
UNRESOLVED_EVIDENCE_LABEL = "referencia interna disponible en Oracle"
EXPORT_NOTICE = (
    "AVISO: este fichero es un borrador declarado. Requiere revisión humana. "
    "No es hecho oficial ni documento presentable."
)
# Public Spanish label for the machine gate code ``draft_requires_human_edit``.
HUMAN_GATE_PUBLIC_LABEL = (
    "Requiere revisión y edición humana antes de presentar. "
    "El borrador no es documento presentable."
)
CHECKLIST_STATUS_LABELS = {
    "pending": "pendiente",
    "ready": "listo",
    "blocked": "bloqueado",
}
GAP_SEVERITY_LABELS = {
    "blocking": "bloqueante",
    "important": "importante",
    "info": "informativo",
}

_UUID_RE = re.compile(r"(?i)\b[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}\b")
# Multi-segment snake_case identifiers (field names / machine codes).
_SNAKE_CASE_RE = re.compile(r"\b[a-z][a-z0-9]*(?:_[a-z0-9]+)+\b")
_CRLF_RE = re.compile(r"[\r\n\x00]+")
_UNSAFE_FILENAME_RE = re.compile(r"[^A-Za-z0-9._\- áéíóúüñÁÉÍÓÚÜÑ]+")

# Explicit internal tokens that must never appear in client-facing DOCX text.
_INTERNAL_TOKEN_DENYLIST: frozenset[str] = frozenset(
    {
        "draft_requires_human_edit",
        "declared_draft",
        "declared_generated",
        "human_gate",
        "our_response_draft",
        "points_hint",
        "requirement_origin",
        "response_origin",
        "official_evidence_ids",
        "declared_evidence_ids",
        "gaps_summary",
        "administrative_checklist",
        "based_on_verdict",
        "draft_engine",
        "tender_ref",
        "lot_hint",
        "source_artifact_id",
        "last_edited_by_user_id",
        "tenant_id",
        "user_id",
        "audit_id",
        "go_conditioned",
        "award_economic",
        "award_technical",
        "sv2_borrador_v1",
        "sv2_prosa_v1",
    }
)


@dataclass(frozen=True, slots=True)
class EvidenceCitation:
    """Human-facing citation resolved inside tenant/dossier scope."""

    title: str
    source: str | None = None
    url: str | None = None
    resolved: bool = True

    def display_line(self) -> str:
        if not self.resolved:
            return UNRESOLVED_EVIDENCE_LABEL
        parts = [self.title]
        if self.source:
            parts.append(f"Fuente: {self.source}")
        if self.url:
            parts.append(f"URL: {self.url}")
        return " · ".join(parts)


def sanitize_export_filename(dossier_title: str, *, version: int) -> str:
    """Safe, legible attachment name derived from the expediente (no paths/CRLF)."""

    base = _CRLF_RE.sub(" ", str(dossier_title or "")).strip()
    base = base.replace("\\", " ").replace("/", " ")
    base = _UNSAFE_FILENAME_RE.sub("-", base)
    base = re.sub(r"\s+", " ", base).strip(" .-_")
    if not base:
        base = "borrador-oferta"
    base = base[:80].rstrip(" .-_")
    ver = max(1, int(version))
    name = f"Borrador-oferta-{base}-v{ver}.docx"
    name = "".join(ch for ch in name if ch.isprintable() and ch not in {"\r", "\n", "\x00"})
    if not name.lower().endswith(".docx"):
        name = f"{name}.docx"
    return name[:180] or "Borrador-oferta.docx"


def content_disposition_attachment(filename: str) -> str:
    """RFC 6266 attachment header; ASCII fallback + UTF-8 filename*."""

    cleaned = "".join(
        ch
        for ch in _CRLF_RE.sub("", str(filename or ""))
        if ch.isprintable() and ch not in {"/", "\\"}
    ).strip()
    if not cleaned.lower().endswith(".docx"):
        cleaned = f"{cleaned or 'Borrador-oferta'}.docx"
    safe = cleaned[:180] or "Borrador-oferta.docx"
    ascii_name = (
        "".join(ch if 32 <= ord(ch) < 127 and ch not in {'"', "\\"} else "_" for ch in safe)
        or "Borrador-oferta.docx"
    )
    starred = quote(safe, safe="")
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{starred}"


def collect_evidence_ids(content: Mapping[str, Any]) -> list[uuid.UUID]:
    """Stable unique evidence UUID list from durable draft content."""

    seen: set[uuid.UUID] = set()
    ordered: list[uuid.UUID] = []

    def _add(raw: Any) -> None:
        if not isinstance(raw, list):
            return
        for item in raw:
            text = str(item or "").strip()
            if not text:
                continue
            try:
                value = uuid.UUID(text)
            except (TypeError, ValueError):
                continue
            if value in seen:
                continue
            seen.add(value)
            ordered.append(value)
            if len(ordered) >= MAX_TRACEABILITY_ROWS:
                return

    _add(content.get("official_evidence_ids"))
    _add(content.get("declared_evidence_ids"))
    for sec in content.get("sections") or []:
        if not isinstance(sec, dict):
            continue
        _add(sec.get("official_evidence_ids"))
        _add(sec.get("declared_evidence_ids"))
        if len(ordered) >= MAX_TRACEABILITY_ROWS:
            break
    return ordered


def resolve_evidence_citations(
    evidence_ids: list[uuid.UUID],
    *,
    lookup: Mapping[uuid.UUID, EvidenceCitation] | None = None,
) -> list[EvidenceCitation]:
    """Map IDs to citations; unresolved IDs become the honest fallback label."""

    resolved_lookup = lookup or {}
    out: list[EvidenceCitation] = []
    for eid in evidence_ids[:MAX_TRACEABILITY_ROWS]:
        hit = resolved_lookup.get(eid)
        if hit is None:
            out.append(
                EvidenceCitation(
                    title=UNRESOLVED_EVIDENCE_LABEL,
                    resolved=False,
                )
            )
        else:
            out.append(hit)
    return out


def assert_no_internal_ids(text: str) -> None:
    """Raise if machine tokens / bare UUIDs / snake_case field names leak into DOCX text."""

    if _UUID_RE.search(text):
        raise OfferDraftError(
            "El export no puede incluir identificadores internos.",
            code="export_sanitization_failed",
            status=500,
        )
    lowered = text.casefold()
    for forbidden in sorted(_INTERNAL_TOKEN_DENYLIST):
        if forbidden.casefold() in lowered:
            raise OfferDraftError(
                "El export no puede incluir códigos o campos internos.",
                code="export_sanitization_failed",
                status=500,
            )
    # Catch any remaining multi-segment snake_case identifiers (machine field names).
    for match in _SNAKE_CASE_RE.finditer(lowered):
        token = match.group(0)
        # Skip pure numeric suffixes already covered; any a_b form is internal.
        raise OfferDraftError(
            f"El export no puede incluir tokens internos ({token}).",
            code="export_sanitization_failed",
            status=500,
        )


def _clean_text(value: Any) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    return "".join(ch if ch in "\t\n" or (ord(ch) >= 32 and ch != "\x7f") else " " for ch in text)


def humanize_human_gate(raw: Any) -> str:
    """Map machine human_gate codes to a public Spanish phrase (never print the code)."""

    gate = str(raw or _HUMAN_GATE).strip() or _HUMAN_GATE
    if gate == _HUMAN_GATE or gate.casefold() == "draft_requires_human_edit":
        return HUMAN_GATE_PUBLIC_LABEL
    # Unknown gate values: still avoid printing snake_case codes.
    if _SNAKE_CASE_RE.search(gate) or gate in _INTERNAL_TOKEN_DENYLIST:
        return HUMAN_GATE_PUBLIC_LABEL
    cleaned = _clean_text(gate)
    if not cleaned or _UUID_RE.search(cleaned):
        return HUMAN_GATE_PUBLIC_LABEL
    return cleaned


def _set_run_shading(run: Any, fill: str = "FFF3CD") -> None:
    """Highlight run background for the export notice (editable, not an image)."""

    r_pr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    r_pr.append(shd)


def _set_paragraph_pagination(
    paragraph: Any,
    *,
    keep_with_next: bool | None = None,
    keep_together: bool | None = None,
    widow_control: bool | None = True,
) -> Any:
    """Apply Word pagination controls (keepNext / keepLines / widowControl)."""

    pf = paragraph.paragraph_format
    if keep_with_next is not None:
        pf.keep_with_next = keep_with_next
    if keep_together is not None:
        pf.keep_together = keep_together
    if widow_control is not None:
        pf.widow_control = widow_control
    return paragraph


def _cohere_block(paragraphs: list[Any]) -> None:
    """Keep a visual section block together when it fits on one page.

    Every paragraph except the last gets ``keep_with_next`` so Word does not
    orphan a lone bullet/gap at the top of the next page. Long sections may
    still break naturally mid-block when they exceed a page.
    """

    if not paragraphs:
        return
    last_idx = len(paragraphs) - 1
    for idx, paragraph in enumerate(paragraphs):
        _set_paragraph_pagination(
            paragraph,
            keep_with_next=idx < last_idx,
            keep_together=True,
            widow_control=True,
        )


def _add_heading(doc: Document, text: str, level: int) -> Any:
    heading = doc.add_heading(_clean_text(text), level=level)
    _set_paragraph_pagination(heading, keep_with_next=True, keep_together=True, widow_control=True)
    return heading


def _add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
) -> Any:
    p = doc.add_paragraph()
    run = p.add_run(_clean_text(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    _set_paragraph_pagination(p, keep_together=True, widow_control=True)
    return p


def _add_list_item(doc: Document, text: str) -> Any:
    p = doc.add_paragraph(_clean_text(text), style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    _set_paragraph_pagination(p, keep_together=True, widow_control=True)
    return p


def _add_notice(doc: Document, text: str) -> Any:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(_clean_text(text))
    run.bold = True
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x83, 0x3C, 0x0C)
    _set_run_shading(run, "FFF3CD")
    _set_paragraph_pagination(p, keep_together=True, widow_control=True)
    return p


def _add_checklist_table(doc: Document, rows: list[list[str]]) -> Any:
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    headers = ("Estado", "Elemento", "Descripción")
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            _set_paragraph_pagination(paragraph, keep_together=True, widow_control=True)
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(3):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = _clean_text(row[c_idx] if c_idx < len(row) else "")
            for paragraph in cell.paragraphs:
                _set_paragraph_pagination(paragraph, keep_together=True, widow_control=True)
    return table


def build_offer_draft_document(
    content: Mapping[str, Any],
    *,
    dossier_title: str,
    version: int,
    exported_at: datetime,
    citations: list[EvidenceCitation] | None = None,
) -> Document:
    """Build a python-docx Document from durable draft content (ordered contract)."""

    doc = Document()
    # Sensible defaults for editable commercial draft.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    banner = str(content.get("banner") or _BANNER).strip() or _BANNER
    statement = str(content.get("statement") or "").strip()
    tender_ref = (
        str(content.get("tender_ref")).strip()
        if content.get("tender_ref") not in (None, "")
        else ""
    )
    lot_hint = (
        str(content.get("lot_hint")).strip() if content.get("lot_hint") not in (None, "") else ""
    )
    export_stamp = exported_at.astimezone(UTC).strftime("%Y-%m-%d %H:%M UTC")

    # 1. Cover / title
    title = doc.add_heading("Borrador de oferta", level=0)
    _set_paragraph_pagination(title, keep_with_next=True, keep_together=True, widow_control=True)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    _add_heading(doc, str(dossier_title or "Expediente").strip() or "Expediente", 1)
    if tender_ref:
        _add_para(doc, f"Referencia: {tender_ref}")
    if lot_hint:
        _add_para(doc, f"Lote: {lot_hint}")
    _add_para(doc, f"Versión del borrador: v{int(version)}")
    _add_para(doc, f"Fecha de exportación: {export_stamp}")
    doc.add_paragraph("")

    # 2. Highlighted notice — public language only (no machine gate codes).
    notice_block: list[Any] = [_add_notice(doc, EXPORT_NOTICE)]
    if banner and banner.casefold() not in EXPORT_NOTICE.casefold():
        p = doc.add_paragraph()
        run = p.add_run(_clean_text(banner))
        run.italic = True
        run.font.size = Pt(11)
        _set_paragraph_pagination(p, keep_together=True, widow_control=True)
        notice_block.append(p)
    # Always surface the human gate in Spanish; never emit draft_requires_human_edit.
    gate_para = _add_para(doc, f"Puerta humana: {humanize_human_gate(content.get('human_gate'))}")
    notice_block.append(gate_para)
    _cohere_block(notice_block)
    doc.add_paragraph("")

    # 3. Introduction
    intro_heading = _add_heading(doc, "Introducción", 1)
    intro_body = _add_para(doc, statement or "(sin introducción)")
    _cohere_block([intro_heading, intro_body])
    doc.add_paragraph("")

    # 4. Sections — each section is a cohesive keep-with-next block.
    _add_heading(doc, "Secciones de la oferta", 1)
    sections = [sec for sec in (content.get("sections") or []) if isinstance(sec, dict)]
    sections = sorted(
        sections,
        key=lambda s: int(s.get("order") if s.get("order") is not None else 0),
    )
    if not sections:
        _add_para(doc, "(sin secciones)")
    for sec in sections:
        # Prefer human title; never fall back to snake_case section keys.
        title_text = str(sec.get("title") or "").strip() or "Sección"
        block: list[Any] = [_add_heading(doc, title_text, 2)]
        requirement = str(sec.get("requirement") or "").strip()
        if requirement:
            # Preserve useful semantic tags like «requisito oficial» / «[oficial]».
            block.append(_add_para(doc, f"Criterio / requisito oficial: {requirement}"))
        points = sec.get("points_hint")
        if points not in (None, ""):
            block.append(_add_para(doc, f"Puntos: {points}"))
        response = str(sec.get("our_response_draft") or "").strip()
        if response:
            lowered = response.casefold()
            if "[borrador declarado" not in lowered and "no es hecho" not in lowered:
                response = f"{_RESPONSE_TAG} {response}"
            block.append(_add_para(doc, f"Respuesta (borrador declarado): {response}"))
        for gap in sec.get("gaps") or []:
            gap_text = str(gap or "").strip()
            if gap_text:
                block.append(_add_list_item(doc, f"Gap: {gap_text}"))
        _cohere_block(block)
    doc.add_paragraph("")

    # 5. Gaps / conditions + administrative checklist
    gaps_heading = _add_heading(doc, "Gaps y condiciones", 1)
    gaps_block: list[Any] = [gaps_heading]
    gaps_summary = [str(g).strip() for g in (content.get("gaps_summary") or []) if str(g).strip()]
    structured_gaps = [
        g
        for g in (content.get("gaps") or [])
        if isinstance(g, dict) and str(g.get("description") or "").strip()
    ]
    if gaps_summary:
        for g in gaps_summary:
            gaps_block.append(_add_list_item(doc, g))
    elif structured_gaps:
        for g in structured_gaps:
            sev_raw = str(g.get("severity") or "important")
            sev = GAP_SEVERITY_LABELS.get(sev_raw, "importante")
            desc = str(g.get("description") or "").strip()
            gaps_block.append(_add_list_item(doc, f"[{sev}] {desc}"))
    else:
        gaps_block.append(_add_para(doc, "(sin gaps registrados)"))
    _cohere_block(gaps_block)

    doc.add_paragraph("")
    checklist_heading = _add_heading(doc, "Checklist administrativa", 1)
    _set_paragraph_pagination(
        checklist_heading, keep_with_next=True, keep_together=True, widow_control=True
    )
    checklist = [
        item for item in (content.get("administrative_checklist") or []) if isinstance(item, dict)
    ]
    if checklist:
        rows: list[list[str]] = []
        for item in checklist:
            status_raw = str(item.get("status") or "pending")
            status = CHECKLIST_STATUS_LABELS.get(status_raw, "pendiente")
            # Prefer human label; never fall back to machine keys like "deuc".
            label = str(item.get("label") or "").strip() or "Elemento"
            desc = str(item.get("description") or "").strip()
            rows.append([status, label, desc])
        _add_checklist_table(doc, rows)
    else:
        _add_para(doc, "(sin checklist administrativa)")
    doc.add_paragraph("")

    # 6. Traceability annex (no bare UUIDs)
    annex_heading = _add_heading(doc, "Anexo de trazabilidad", 1)
    annex_block: list[Any] = [
        annex_heading,
        _add_para(
            doc,
            "Referencias de evidencia resueltas dentro del expediente. "
            "No se imprimen identificadores internos.",
        ),
    ]
    cites = citations if citations is not None else []
    if cites:
        for idx, cite in enumerate(cites, start=1):
            annex_block.append(_add_list_item(doc, f"{idx}. {cite.display_line()}"))
    else:
        annex_block.append(_add_para(doc, "(sin referencias de evidencia en el borrador)"))
    _cohere_block(annex_block)

    doc.add_paragraph("")
    _add_para(
        doc,
        "Nota final: este texto es un borrador declarado — "
        "no es hecho oficial ni documento presentable.",
    )

    # Soft paragraph cap (document body children).
    body_count = len(doc.element.body)
    if body_count > MAX_DOCX_PARAGRAPHS:
        raise OfferDraftError(
            "El documento supera el número máximo de párrafos exportables.",
            code="payload_too_large",
            status=422,
        )
    return doc


def build_offer_draft_docx(
    content: Mapping[str, Any],
    *,
    dossier_title: str,
    version: int,
    exported_at: datetime | None = None,
    citations: list[EvidenceCitation] | None = None,
) -> bytes:
    """Render the durable draft as a complete DOCX package (bytes)."""

    if not isinstance(content, Mapping) or not content:
        raise OfferDraftError(
            "No hay contenido de borrador para exportar.",
            code="offer_draft_not_found",
            status=404,
        )
    raw_probe = str(content.get("statement") or "")
    if len(raw_probe.encode("utf-8")) > MAX_CONTENT_BYTES:
        raise OfferDraftError(
            "El borrador supera el tamaño máximo permitido.",
            code="payload_too_large",
            status=422,
        )

    when = exported_at or datetime.now(UTC)
    doc = build_offer_draft_document(
        content,
        dossier_title=dossier_title,
        version=version,
        exported_at=when,
        citations=citations,
    )

    # Collect plain text for sanitization gate.
    plain_parts: list[str] = []
    for para in doc.paragraphs:
        plain_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                plain_parts.append(cell.text)
    plain = "\n".join(plain_parts)
    assert_no_internal_ids(plain)

    buf = io.BytesIO()
    doc.save(buf)
    payload = buf.getvalue()
    if len(payload) > MAX_DOCX_BYTES:
        raise OfferDraftError(
            "El documento exportado supera el límite de tamaño.",
            code="payload_too_large",
            status=422,
        )
    if not payload.startswith(b"PK"):
        raise OfferDraftError(
            "Fallo al empaquetar el documento Word.",
            code="export_failed",
            status=500,
        )
    return payload


def load_evidence_citation_lookup(
    *,
    tenant_id: uuid.UUID,
    dossier_id: uuid.UUID,
    evidence_ids: list[uuid.UUID],
) -> dict[uuid.UUID, EvidenceCitation]:
    """Resolve evidence titles/sources/URLs within tenant + dossier linkage.

    Unlinked or cross-tenant IDs are omitted (caller falls back to honest label).
    """

    if not evidence_ids:
        return {}

    from sqlalchemy import select

    from opn_oracle.documents.models import Document
    from opn_oracle.extensions import db
    from opn_oracle.oracle.links import EvidenceDossier
    from opn_oracle.oracle.models import Evidence, Signal

    linked_ids = set(
        db.session.scalars(
            select(EvidenceDossier.evidence_id).where(
                EvidenceDossier.tenant_id == tenant_id,
                EvidenceDossier.dossier_id == dossier_id,
                EvidenceDossier.evidence_id.in_(evidence_ids),
            )
        )
    )
    if not linked_ids:
        return {}

    rows = list(
        db.session.scalars(
            select(Evidence).where(
                Evidence.tenant_id == tenant_id,
                Evidence.id.in_(list(linked_ids)),
            )
        )
    )
    signal_ids = [row.signal_id for row in rows if row.signal_id is not None]
    document_ids = [row.document_id for row in rows if row.document_id is not None]
    signals: dict[uuid.UUID, Signal] = {}
    if signal_ids:
        for sig in db.session.scalars(
            select(Signal).where(Signal.tenant_id == tenant_id, Signal.id.in_(signal_ids))
        ):
            signals[sig.id] = sig
    documents: dict[uuid.UUID, Document] = {}
    if document_ids:
        for doc_row in db.session.scalars(
            select(Document).where(Document.tenant_id == tenant_id, Document.id.in_(document_ids))
        ):
            documents[doc_row.id] = doc_row

    out: dict[uuid.UUID, EvidenceCitation] = {}
    for row in rows:
        title: str | None = None
        source: str | None = None
        url = (str(row.source_url).strip() if row.source_url else None) or None
        provenance = row.provenance if isinstance(row.provenance, dict) else {}

        if row.signal_id and row.signal_id in signals:
            sig = signals[row.signal_id]
            title = str(sig.title or "").strip() or None
            source = str(sig.source_name or sig.source_type or "").strip() or None
            if not url and sig.source_url:
                url = str(sig.source_url).strip() or None
        elif row.document_id and row.document_id in documents:
            doc_row = documents[row.document_id]
            title = str(doc_row.original_filename or "").strip() or "Documento del expediente"
            source = "documento"
        else:
            for key in ("title", "label", "name", "source_title"):
                raw = provenance.get(key)
                if raw:
                    title = str(raw).strip()[:300] or None
                    if title:
                        break
            for key in ("source_name", "source", "provider", "source_kind"):
                raw = provenance.get(key)
                if raw:
                    source = str(raw).strip()[:240] or None
                    if source:
                        break
            if not source:
                source = str(row.source_kind or "").strip() or None
            if not url:
                for key in ("source_url", "url", "canonical_url"):
                    raw = provenance.get(key)
                    if raw:
                        url = str(raw).strip()[:1500] or None
                        if url:
                            break

        if not title:
            extract = str(row.extract or "").strip()
            if extract:
                title = (extract[:120] + "…") if len(extract) > 120 else extract
            else:
                continue

        if _UUID_RE.search(title) or (source and _UUID_RE.search(source)):
            continue
        if url and not (url.startswith("http://") or url.startswith("https://")):
            url = None

        out[row.id] = EvidenceCitation(
            title=title[:300],
            source=source[:240] if source else None,
            url=url[:1500] if url else None,
            resolved=True,
        )
    return out


__all__ = [
    "DOCX_MEDIA_TYPE",
    "EXPORT_NOTICE",
    "HUMAN_GATE_PUBLIC_LABEL",
    "MAX_DOCX_BYTES",
    "UNRESOLVED_EVIDENCE_LABEL",
    "EvidenceCitation",
    "assert_no_internal_ids",
    "build_offer_draft_document",
    "build_offer_draft_docx",
    "collect_evidence_ids",
    "content_disposition_attachment",
    "humanize_human_gate",
    "load_evidence_citation_lookup",
    "resolve_evidence_citations",
    "sanitize_export_filename",
    "xml_escape",
]
